# -*- coding: utf-8 -*-
from __future__ import print_function
from requests import get
from bs4 import BeautifulSoup
import docx
import os
import pickle
import os.path
import sys
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.http import MediaFileUpload

SCOPES = ['https://www.googleapis.com/auth/drive']
niepotrzebne = ['© Copyright 2018 - Zespół Szkół Łączności w Gdańsku','Podwale Staromiejskie 51/52, 80-845 Gdańsk','58 301 13 77','58 308 01 00']
errory = []

def creds(): # authenticate the user
    creds = None
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token, encoding="bytes")
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server()
        # Save the credentials for the next run
        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)
    return creds

def raw(url): # get and parse the article
    response = get(url)
    html_soup = BeautifulSoup(response.text, 'html.parser')
    akt = html_soup.find_all('div', class_='nicdark_small_news')
    return akt

def tworze(blok, tytul, z): # create the document on disc
    doc = docx.Document('wzor.docx')
    elo = blok.div.div.h3
    naglowek = elo.text
    print(naglowek)
    doc.add_paragraph(naglowek).style = 'kron'
    for a in elo.find_all('a', href=True):
        link = a['href']
        url_akt = 'http://www.zsl.gda.pl' + link
        odp = get(url_akt)
        zupa = BeautifulSoup(odp.text, 'html.parser')
        div2 = zupa.find_all('p')
        for i in div2:
            if i.text not in niepotrzebne:
                zaw = i.text
                doc.add_paragraph(zaw).style = 'krot'
    for src in blok.div.div.find_all('img', class_='nicdark_section', src=True):
        tmp = 'http://www.zsl.gda.pl' + src['src']
        obrazek(tmp)
        doc.add_picture('tmp.png')
    #zapisuje plik z uwzglednieniem mozliwych bledow, zapisuje je do error loga
    try:
        doc.save('C:\\kronika\\' + tytul + '.docx')
    except:
        doc.save('C:\\kronika\\%s.docx' % z)
        errory.append(tytul)

def obrazek(tmp): # create the image
    png = get(tmp)
    with open('tmp.png', 'wb') as f:
        f.write(png.content)
    return

def dzialanie(x, z):
    for i in range(x, -1, -9):
        url = 'http://www.zsl.gda.pl/aktualnosci/?start=%s' % i
        akt = raw(url)
        for blok in reversed(akt):
            #sprawdzam czy plik istnieje
            elo = blok.div.div.h3
            naglowek_tmp = elo.text
            tytul = str(z) + ' - ' + naglowek_tmp
            tytul = str(tytul)
            if (os.path.isfile('C:\\kronika\\%s.docx' % tytul)):
                print('plik juz istnieje')
            else:
            #tworze dokument i uzupełniam treść
                tworze(blok, tytul, z)
            z = z + 1
        print(i)

def main():
    crede = creds()
    z = 0
    service = build('drive', 'v3', credentials=crede)
    results = service.files().list(
        pageSize=10, fields="nextPageToken, files(id, name)").execute()
    pliki = results.get('files', [])
    
    if not os.path.exists('var.txt'):
        print("Sprawdzam ilosc newsow...")
        print("To moze chwile potrwac...")
        i = 0
        while z == 0:
            url = 'http://www.zsl.gda.pl/aktualnosci/?start=%s' % i
            akt = raw(url)
            i += 9
            if akt == []:
                z = 1
                with open('var.txt','w') as g:
                    g.write(str(i))
    with open('var.txt', 'r') as r:
        i = r.read()
    i = int(i)
    if not os.path.exists('C:\\kronika'):
        os.makedirs('C:\\kronika')
    dzialanie(i, z)
    for file in os.listdir('C:\\kronika'):
        if file not in pliki:
            file_metadata = {'name': file}
            media = MediaFileUpload(str(file),mimetype='text/docx')
            file = service.files().create(body=file_metadata,media_body=media,fields='id').execute()
    with open('C:\\kronika\\errorlog.txt','w') as f:
        for blad in errory:
            f.write('%s\n' % blad)
    print('Lista nazw plikow ktore powoduja errory:')
    print("\n".join(errory))
    print('\nLista plikow znajduje sie tez w errorlog.txt\n')
    input("Wcisnij dowolny przycisk zeby zakonczyc")
main()
